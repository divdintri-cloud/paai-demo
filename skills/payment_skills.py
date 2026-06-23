from datetime import date
from io import BytesIO

import pandas as pd
from docx import Document

from tools.openai_client import call_model


PAYMENT_FILE = "data/payment_reminders.csv"


def load_payment_reminders():
    try:
        return pd.read_csv(PAYMENT_FILE)
    except FileNotFoundError:
        return pd.DataFrame(
            columns=[
                "Payment",
                "Category",
                "Due Date",
                "Amount",
                "Status",
                "Reminder Days Before",
                "Notes",
            ]
        )


def save_payment_reminders(new_df):
    existing_df = load_payment_reminders()

    if existing_df.empty:
        combined_df = new_df.copy()
    else:
        combined_df = pd.concat([existing_df, new_df], ignore_index=True)

    if "Payment" in combined_df.columns and "Due Date" in combined_df.columns:
        combined_df["Payment"] = combined_df["Payment"].fillna("").astype(str).str.strip()
        combined_df["Due Date"] = combined_df["Due Date"].fillna("").astype(str).str.strip()

        combined_df = combined_df.drop_duplicates(
            subset=["Payment", "Due Date"],
            keep="first",
        )

    combined_df.to_csv(PAYMENT_FILE, index=False)
    return combined_df


def get_payment_summary():
    df = load_payment_reminders()

    if df.empty:
        return {
            "dataframe": df,
            "overdue": df,
            "due_soon": df,
            "subscriptions": df,
            "summary": {
                "total_payments": 0,
                "overdue_count": 0,
                "due_soon_count": 0,
                "subscriptions_count": 0,
                "total_pending_amount": 0,
            },
        }

    df["Due Date"] = pd.to_datetime(df["Due Date"], errors="coerce")
    today = pd.Timestamp(date.today())

    df["Days Until Due"] = (df["Due Date"] - today).dt.days

    unpaid_df = df[df["Status"].astype(str).str.lower() != "paid"]

    overdue_df = unpaid_df[unpaid_df["Days Until Due"] < 0]

    due_soon_df = unpaid_df[
        (unpaid_df["Days Until Due"] >= 0)
        & (unpaid_df["Days Until Due"] <= 7)
    ]

    subscriptions_df = df[
        df["Category"].astype(str).str.lower().str.contains("subscription", na=False)
    ]

    total_pending_amount = unpaid_df["Amount"].fillna(0).sum()

    summary = {
        "total_payments": len(df),
        "overdue_count": len(overdue_df),
        "due_soon_count": len(due_soon_df),
        "subscriptions_count": len(subscriptions_df),
        "total_pending_amount": total_pending_amount,
    }

    return {
        "dataframe": df,
        "overdue": overdue_df,
        "due_soon": due_soon_df,
        "subscriptions": subscriptions_df,
        "summary": summary,
    }


def read_excel_or_csv(uploaded_file):
    file_name = uploaded_file.name.lower()

    if file_name.endswith(".csv"):
        return pd.read_csv(uploaded_file)

    if file_name.endswith(".xlsx") or file_name.endswith(".xls"):
        return pd.read_excel(uploaded_file)

    return pd.DataFrame()


def read_word_document(uploaded_file):
    document = Document(BytesIO(uploaded_file.read()))
    paragraphs = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs)


def standardize_payment_dataframe(df):
    """
    Converts common uploaded expense/payment columns into PAAI payment reminder format.
    This is a best-effort mapper. User should review before saving.
    """

    if df.empty:
        return pd.DataFrame(
            columns=[
                "Payment",
                "Category",
                "Due Date",
                "Amount",
                "Status",
                "Reminder Days Before",
                "Notes",
            ]
        )

    normalized_columns = {col.lower().strip(): col for col in df.columns}

    def find_column(possible_names):
        for name in possible_names:
            if name in normalized_columns:
                return normalized_columns[name]
        return None

    payment_col = find_column(["payment", "merchant", "vendor", "name", "description", "bill"])
    category_col = find_column(["category", "type"])
    due_date_col = find_column(["due date", "date", "payment date", "bill date"])
    amount_col = find_column(["amount", "cost", "payment amount", "total"])
    status_col = find_column(["status", "paid status"])
    notes_col = find_column(["notes", "memo", "comments"])

    output_df = pd.DataFrame()

    output_df["Payment"] = df[payment_col] if payment_col else "Unclear"
    output_df["Category"] = df[category_col] if category_col else "Uncategorized"
    output_df["Due Date"] = df[due_date_col] if due_date_col else ""
    output_df["Amount"] = df[amount_col] if amount_col else 0
    output_df["Status"] = df[status_col] if status_col else "Pending"
    output_df["Reminder Days Before"] = 3
    output_df["Notes"] = df[notes_col] if notes_col else "Imported from uploaded file. Review before saving."

    return output_df


def extract_payments_from_word_text(text):
    system_prompt = """
You are the Payment Reminder Agent for PAAI.

The user uploaded a Word document that may contain payment reminders, bills, subscriptions, or due dates.

Extract payment reminder rows.

Return only valid CSV text with these exact columns:
Payment,Category,Due Date,Amount,Status,Reminder Days Before,Notes

Rules:
- Do not invent payments.
- If amount is missing, use 0.
- If status is missing, use Pending.
- If reminder days are missing, use 3.
- If due date is missing, leave it blank.
- Do not include markdown.
"""

    user_prompt = f"""
Document text:
{text}
"""

    csv_text = call_model(system_prompt, user_prompt)

    try:
        return pd.read_csv(BytesIO(csv_text.encode("utf-8")))
    except Exception:
        return pd.DataFrame(
            columns=[
                "Payment",
                "Category",
                "Due Date",
                "Amount",
                "Status",
                "Reminder Days Before",
                "Notes",
            ]
        )


def parse_uploaded_payment_file(uploaded_file):
    file_name = uploaded_file.name.lower()

    if file_name.endswith(".csv") or file_name.endswith(".xlsx") or file_name.endswith(".xls"):
        raw_df = read_excel_or_csv(uploaded_file)
        standardized_df = standardize_payment_dataframe(raw_df)
        return standardized_df

    if file_name.endswith(".docx"):
        text = read_word_document(uploaded_file)
        return extract_payments_from_word_text(text)

    return pd.DataFrame()


def ask_payment_agent(question):
    payment_data = get_payment_summary()
    df = payment_data["dataframe"]

    system_prompt = """
You are the Payment Reminder Agent for PAAI.

Your job is to help the user track bills, subscriptions, renewals, and payment reminders.

Rules:
- Do not make payments.
- Do not ask for bank passwords.
- Do not claim a payment is complete unless the data says it is paid.
- Clearly identify overdue, due soon, and upcoming payments.
- Recommend manual verification before making any financial decision.
- Keep the answer practical and concise.
"""

    user_prompt = f"""
User question:
{question}

Payment reminder data:
{df.to_string(index=False)}

Today's date:
{date.today()}

Answer the user's question.
"""

    return call_model(system_prompt, user_prompt)
